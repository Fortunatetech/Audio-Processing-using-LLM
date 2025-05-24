import streamlit as st
import pandas as pd
import json
import google.generativeai as genai
import os 
from dotenv import load_dotenv

from io import BytesIO
from fpdf import FPDF
from docx import Document
import tempfile

load_dotenv()

# ─────────────────────────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────────────────────────
genai.configure(api_key=st.secrets['GEMINI_API_KEY'])
model = genai.GenerativeModel('models/gemini-1.5-flash-002')

# ─────────────────────────────────────────────────────────────────────────────
# Authentication setup
# ─────────────────────────────────────────────────────────────────────────────
APP_PASSWORD = os.getenv("APP_PASSWORD")

def login(password_input):
    """Simple function to verify the password."""
    if password_input == APP_PASSWORD:
        st.session_state.logged_in = True
        st.success("🔓 Login successful! You can now use the app.")
    else:
        st.error("❌ Incorrect password. Please try again.")

# Initialize login state
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'report_text' not in st.session_state:
    st.session_state.report_text = None

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("🔒 Login")
    pwd = st.text_input("Application password:", type="password")
    if st.button("Login"):
        login(pwd)

st.sidebar.markdown("---")
st.sidebar.header("📁 Data Upload")
uploaded_file = st.sidebar.file_uploader(
    "Excel file (xlsx)", type=["xlsx"]
)

st.sidebar.markdown("---")
st.sidebar.header("📘 Guidelines")
st.sidebar.markdown("""
- Upload the BNA Excel
- Select the sheet
- Preview Federal vs State data
- Ask a focused question
""")

# st.sidebar.markdown("---")
# st.sidebar.header("📅 Date")
# # default to today (but user can pick any 2025 date)
# _ = st.sidebar.date_input(
#     "Select a date", 
#     value=date(2025, 1, 1), 
#     min_value=date(2025, 1, 1), 
#     max_value=date(2025, 12, 31)
# )

# ─────────────────────────────────────────────────────────────────────────────
# MAIN AREA
# ─────────────────────────────────────────────────────────────────────────────
st.title("📊 Wage Payment Law Analyzer")

if not st.session_state.logged_in:
    st.info("Please log in via the sidebar to proceed.")
    st.stop()

# Logged-in users see the rest:
if uploaded_file:
    try:
        xlsx = pd.ExcelFile(uploaded_file)
        sheets = xlsx.sheet_names
        selected_sheet = st.selectbox("Select a sheet to preview", sheets)
        df = pd.read_excel(xlsx, sheet_name=selected_sheet)

        st.subheader(f"First 5 rows of “{selected_sheet}”")
        st.dataframe(df.head())

        # Federal overview
        first_col = df.columns[0]
        fed = df[df[first_col].str.contains("Federal", case=False, na=False)]
        if not fed.empty:
            st.subheader("📌 Overview of Federal Guidelines")
            st.write(fed)

        # Centered Q&A
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 100, 1])
        with col2:
            user_q = st.text_area("🔍 Ask about deviations vs. Federal law")
            if st.button("Generate Report"):
                if not user_q.strip():
                    st.error("Please type a question before generating.")
                else:
                    with st.spinner("Analyzing…"):
                        # prepare prompt
                        df_json = df.to_dict(orient="records")
                        PROMPT = f"""
                                    You are an expert legal research assistant specializing in U.S. employment law.
                                    Your task is to analyze the provided dataset containing federal guidelines and state-specific laws across multiple categories (Frequency of Payment Analysis, Method of Payment Analysis, Pay Statements Analysis, Notification Requirements Analysis, etc.).
                                    For any category the user mention talk about, 
                                     1. Identify and succinctly summarize the Federal guideline.
                                     2. Compare every state’s law to that Federal guideline, and classify each state into one of four buckets:
                                        • States(names) that align exactly with the Federal guideline.
                                        • States(names) that align but include additional requirements (describe those additions).
                                        • States(names) that deviate from the Federal guideline (describe the deviations).
                                        • States(names) with no specific provision on this category.
                                    3.  Then provide a consolidated summary report that includes:
                                        - The count of states in each bucket.
                                        - Notable patterns or regional trends.
                                        - Any ambiguities or interpretation notes where state law is unclear.
                                    4. If the question isn't relevant, say you don’t know.
                                    5. Note: Ensure all your responses are short and coincides. 
                                    User’s question:
                                    {user_q}

                                    Data:
                                    {json.dumps(df_json, indent=2)}
                                    """
                                # call Gemini
                        try:
                            resp = model.generate_content(PROMPT, request_options={"timeout":300})
                            st.session_state.report_text = resp.text
                        except Exception as e:
                            st.error(f"LLM error: {e}")

                            
    # After generation
        if st.session_state.report_text:
            st.markdown("---")
            st.subheader("📝 Summary Report")
            st.write(st.session_state.report_text)

            # Prepare Word
            doc = Document()
            doc.add_heading('Summary Report', level=1)
            for line in st.session_state.report_text.split('\n'):
                if line.strip() == '':
                    continue
                if line.startswith(('1.','2.','3.','4.','5.')):
                    doc.add_paragraph(line, style='List Number')
                elif line.endswith(':'):
                    doc.add_heading(line.strip(), level=2)
                else:
                    doc.add_paragraph(line)
            word_io = BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            # Prepare PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(True, 15)
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, 'Summary Report', ln=True)
            pdf.set_font('Arial', '', 12)
            for line in st.session_state.report_text.split('\n'):
                if not line.strip():
                    continue
                if line.endswith(':'):
                    pdf.set_font('Arial','B',12)
                    pdf.cell(0, 8, line.strip(), ln=True)
                    pdf.set_font('Arial','',12)
                else:
                    pdf.multi_cell(0, 8, line)
            pdf_bytes = pdf.output(dest='S').encode('latin-1')

            # Download buttons
            c1, c2 = st.columns(2)
            with c1:
                st.download_button('📄 Download Report in PDF Format', data=pdf_bytes, file_name='report.pdf', mime='application/pdf', key='pdf')
            with c2:
                st.download_button('📝 Download Report in Word Format', data=word_io.getvalue(), file_name='report.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key='word')

    except Exception as e:
        st.error(f"Error processing file: {e}")
else:
    st.info("Upload an Excel file in the sidebar to get started.")













































# import streamlit as st
# import pandas as pd
# import json
# import google.generativeai as genai

# # Configure Gemini API
# genai.configure(api_key=st.secrets['GEMINI_API_KEY'])
# model = genai.GenerativeModel('models/gemini-2.0-flash-thinking-exp')

# st.title("📊 Wage Payment Law Analyzer")

# # File uploader
# uploaded_file = st.file_uploader("Upload the BNA Wage Payment Requirements Excel file", type=["xlsx"])

# if uploaded_file:
#     try:
#         # Load the Excel file
#         excel_file = pd.ExcelFile(uploaded_file)

#         # Display available sheet names
#         sheet_names = excel_file.sheet_names
#         #st.write("Available sheets:", sheet_names)

#         # Allow user to select a sheet
#         selected_sheet = st.selectbox("Select a sheet to preview", sheet_names)

#         # Read the selected sheet into a DataFrame
#         df = pd.read_excel(excel_file, sheet_name=selected_sheet)

#         # Display the first 5 rows
#         st.subheader(f"First 5 rows of '{selected_sheet}' sheet:")
#         st.dataframe(df.head())

#         # Identify the column containing 'Federal' information
#         first_col = df.columns[0]
#         federal_row = df[df[first_col].str.contains("Federal", case=False, na=False)]

#         if not federal_row.empty:
#             st.subheader("📌 Overview of Federal Guidelines:")
#             st.write(federal_row)

#         # Text input for user question
#         st.subheader("🔍 Ask a question about the data:")
#         user_question = st.text_area("Enter your question here")

#         if st.button("Submit") and user_question:
#             with st.spinner("Generating response..."):
#                 # Convert DataFrame to JSON
#                 df_json = df.to_dict(orient='records')

#                 # Construct the prompt
#                 PROMPT = f"""
#                 You are an expert legal research assistant specializing in U.S. employment law.

#                 Your task is to analyze the provided dataset containing federal guidelines and state-specific laws across multiple categories (Frequency of Payment Analysis, Method of Payment Analysis, Pay Statements Analysis, Notification Requirements Analysis, etc.).

#                 For each category:
#                 1. Identify and succinctly summarize the Federal guideline.
#                 2. Compare every state’s law to that Federal guideline, and classify each state into one of four buckets:
#                     • States that align exactly with the Federal guideline.
#                     • States that align but include additional requirements (describe those additions).
#                     • States that deviate from the Federal guideline (describe the deviations).
#                     • States with no specific provision on this category.
#                 3. List the names of all states in each bucket.

#                 Then provide a consolidated summary report that includes:
#                 - The count of states in each bucket.
#                 - Notable patterns or regional trends.
#                 - Any ambiguities or interpretation notes where state law is unclear.

#                 Use clear, concise language and structure your output with headings and bullet points.

#                 User’s question:  
#                 {user_question}

#                 Data (JSON array of rows):  
#                 {json.dumps(df_json, indent=2)}
                
#                 """


#                 try:
#                     # Generate response from Gemini
#                     response = model.generate_content(PROMPT)
#                     st.subheader("📝 Summary Response:")
#                     st.write(response.text)
#                 except Exception as e:
#                     st.error(f"An error occurred while generating the response: {e}")
#     except Exception as e:
#         st.error(f"An error occurred while processing the file: {e}")
# else:
#     st.info("Please upload an Excel file to proceed.")
